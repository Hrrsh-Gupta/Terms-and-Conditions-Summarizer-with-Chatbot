import sys
import os
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QTextEdit, QPushButton, QToolButton, QFileDialog, QStackedWidget, QSpacerItem, QSizePolicy, QScrollArea,QFrame, QMenu, QAction, QLineEdit, 
)
from PyQt5.QtGui import QIcon , QPalette, QColor
from PyQt5.QtCore import Qt, QTimer,  QSize, QProcess, QCoreApplication, QEvent
from PyQt5.QtWidgets import QDockWidget

from docx import Document
import fitz
from easyocr import Reader

from keyword_summary import ag_generate_summary
from tfidf_summary import st_generate_summary
from api_summary import generate_api_summary
from database import Database
from groq import Groq


from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors

import warnings
warnings.filterwarnings("ignore", category=DeprecationWarning)


class TermsSummarizerApp(QMainWindow):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("Terms and Conditions Summarizer with Chatbot")
        self.setWindowIcon(QIcon(self.resource_path("./icons/app_icon.png")))
        self.setGeometry(350, 135, 1200, 800)

        self.db = Database()
        self.session_id = self.db.add_log_entry()
        
        
        # Apply dark theme
        self.setStyleSheet("""
            QWidget {
                background-color: #2b2b2b;
                color: #f1f1f1;
            }
            QLabel#titleLabel {
                font-size: 24px;
                font-weight: bold;
                color: white;
            }
            QToolButton, QPushButton {
                background-color: #3c3f41;
                color: #ffffff;
                border: 1px solid #5a5a5a;
                padding: 6px;
                border-radius: 12px;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #505354;
            }
            QTextEdit {
                background-color: #212121;
                color: #ffffff;
                border: 2px solid #5a5a5a;
                border-radius: 6px; 
            }
                           
            QScrollArea {
                background-color: #222222;  /* Sidebar background */
                border: none;  /* Remove border */
            }
            QScrollBar:vertical {
                border: none;
                background-color: #2b2b2b;  /* Dark background for scrollbar */
                width: 13px;
                margin: 0px 0 0px 0;
                border-radius: 0px;
            }
            QScrollBar::handle:vertical {
                background-color: #3f3f3f;  /* Scroll handle color */
                min-height: 20px;
                border-radius: 0px;
            }
            QScrollBar::handle:vertical:hover {
                background-color: #4f4f4f;  /* Hover color for scroll handle */
            }
            
            QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical {
                background-color:#222222; 
            }
                           
        """)

        self.navbar_button = QToolButton()
        self.navbar_button.setCursor(Qt.PointingHandCursor)
        self.navbar_button.setIcon(QIcon(self.resource_path("./icons/menu_icon.png")))  # Replace with path to your menu icon
        self.navbar_button.setIconSize(QSize(24, 24))
        self.navbar_button.clicked.connect(self.toggle_navbar)
        self.navbar_button.setStyleSheet("""
            QToolButton {
                background-color: #2b2b2b;
                border: none;
                padding: 5px;
                border-radius: 6px;
            }
            QToolButton:hover {
                background-color: #505354;
            }
        """)

        self.close_button = QToolButton()
        self.close_button.setCursor(Qt.PointingHandCursor)
        self.close_button.setText("✕")
        self.close_button.setStyleSheet("""
            QToolButton {
                border: none;
                background-color: none;
                font-weight: bold;
                font-size : 16px;
            }
            QToolButton:hover {
                color: red;
            }
        """)
        self.close_button.clicked.connect(lambda: QApplication.quit())
        # close_button.setVisible(self.isMaximized())

        title_bar_layout = QHBoxLayout()
        title_bar_layout.addWidget(self.navbar_button)
        title_bar_layout.addStretch(1)
        title_bar_layout.addWidget(self.close_button)
        
        # Main content setup
        self.stacked_widget = QStackedWidget()
        self.summarizer_page = self.create_summarizer_page()
        self.chatbot_page = self.create_chatbot_page()
        self.stacked_widget.addWidget(self.summarizer_page)
        self.stacked_widget.addWidget(self.chatbot_page)


        # Layout setup
        main_layout = QVBoxLayout()
        main_layout.addLayout(title_bar_layout)
        self.setup_dock_sidebar()
        main_layout.addWidget(self.stacked_widget)

        central_widget = QWidget()
        central_widget.setLayout(main_layout)
        self.setCentralWidget(central_widget)
        self.selected_button = None
        self.prev_session = None



    def setup_dock_sidebar(self):
        self.dock_sidebar = QDockWidget("", self)
        self.dock_sidebar.setAllowedAreas(Qt.LeftDockWidgetArea)
        self.dock_sidebar.setFeatures(QDockWidget.NoDockWidgetFeatures)  # Remove header and close button
        self.dock_sidebar.setVisible(False)  # Sidebar is closed by default
        self.dock_sidebar.setMinimumWidth(160)

        # Dictionary to keep track of buttons and associated page indices
        self.nav_buttons = {}  
        self.total_buttons = 0

        # Scroll area to contain the sidebar
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        

        # Sidebar layout
        sidebar_widget = QWidget()
        self.sidebar_layout = QVBoxLayout()
        self.sidebar_layout.setContentsMargins(10, 5, 10, 0)
        self.sidebar_layout.setSpacing(0)
        self.sidebar_layout.setAlignment(Qt.AlignTop)
        sidebar_widget.setLayout(self.sidebar_layout)

        sidebar_widget.setStyleSheet("""
            QWidget {
                background-color: #222222;
                color: #ffffff;
            }
            QToolButton {
                padding: 14px;
                text-align: left;
                border: none;
                font-size: 16px;
            }
            QToolButton:hover {
                background-color: #373737;
            }
        """)

        # Add buttons dynamically in a loop
        button_names = ["Summarizer", "Chatbot", "Previous Sessions"]
        # session_aliases = []
        sessions= self.db.get_log_entries()
        for session in sessions:
                button_names.append(f"{session[1]} {session[2]}")
                # session_aliases.append(session[3])

        for i, name in enumerate(button_names):
            button = QToolButton()
            if i < 3:
                button.setText(name)
                # if name == 'Chatbot':
                #     button.setVisible(False)
            else:
                but_name = name.split()
                # print(but_name)
                tim_name = but_name[1].split(":")
                but_name = but_name[0] + " " + tim_name[0] + ":" +tim_name[1] + " " + but_name[2]
                button.setText(but_name)

            button.setCursor(Qt.PointingHandCursor)
            
            # Make the "Previous Sessions:" button look like a label
            if name == "Previous Sessions":
                button.setEnabled(False)  # Disable click functionality
                button.setCursor(Qt.ArrowCursor)  # Set cursor to arrow instead of hand
                button.setStyleSheet(""" 
                    QToolButton { 
                    border: none;         
                    color: #878787;                         
                    padding: 10px;
                    padding-left: 15px;
                    padding-bottom: 0px;
                    padding-top: 24px;
                                     
                    text-align: left;
                    border: none;
                    font-size: 14px;
                    font-weight : bold;
                }""")  # Customize appearance
            
            else:
                # Enable other buttons to be clickable
                if i == 0 or i == 1:
                    button.clicked.connect(lambda checked, idx=i: self.show_page(idx))
                else:
                    session_id = sessions[i-3][0]
                    session_name = f"{sessions[i-3][1]} {sessions[i-3][2]}"
                    button.clicked.connect(lambda checked, idx=(session_id,session_name,button): self.show_session(idx))
            
            # Store button and page index in nav_buttons dictionary
            self.nav_buttons[name] = button
            self.sidebar_layout.addWidget(button)
            self.total_buttons += 1

        # self.sidebar_layout.addStretch()
        # bottom_spacer = QSpacerItem(20, 40, QSizePolicy.Minimum, QSizePolicy.Expanding)
        # self.sidebar_layout.addItem(bottom_spacer)

        # Set the sidebar_widget as the scroll area's widget
        scroll_area.setWidget(sidebar_widget)
        self.dock_sidebar.setWidget(scroll_area)
        self.addDockWidget(Qt.LeftDockWidgetArea, self.dock_sidebar)

        self.resizeDocks([self.dock_sidebar], [215], Qt.Horizontal)


    def update_sidebar_sessions(self):
    # Clear existing session buttons

        sessions = self.db.get_log_entries()
        curr_session = self.db.get_log_entry(self.session_id)
        curr_session_name = f"{curr_session[1]} {curr_session[2]}"

        for name, button in list(self.nav_buttons.items()):
            if name not in ["Summarizer", "Chatbot", "Previous Sessions"]:
                
                # if name == curr_session_name:
                #     break
                button.setParent(None)
                del self.nav_buttons[name]
                del button
                self.total_buttons -= 1

        else:   
        
            for session in sessions:
                session_name = f"{session[1]} {session[2]}"

            
                button = QToolButton()
                but_name = session_name.split()
                tim_name = but_name[1].split(":")
                but_name = but_name[0] + " " + tim_name[0] + ":" +tim_name[1] + " " + but_name[2]

                button.setText(but_name)

                # button.setText(session_name)
                button.setCursor(Qt.PointingHandCursor)

                if session_name == curr_session_name:
                    button.setStyleSheet(""" 
                    QToolButton {   
                    background-color: #303030;
                    }              
                    QToolButton:hover {
                        background-color: #373737;
                }""")
                    
                session_id = session[0]
                button.clicked.connect(lambda checked, idx=(session_id,session_name,button): self.show_session(idx))

                # Add the button to the sidebar layout and nav_buttons dictionary
                self.nav_buttons[session_name] = button
                self.sidebar_layout.addWidget(button)  # Use sidebar_layout directly
                self.total_buttons += 1
            # bottom_spacer = QSpacerItem(20, 40, QSizePolicy.Maximum, QSizePolicy.Expanding)
            # self.sidebar_layout.addItem(bottom_spacer)


    def show_page(self, index): 
        if self.selected_button:
            if self.prev_session == self.session_id:
                    # print("yes")
                    self.selected_button.setStyleSheet(""" 
                    QToolButton {   
                        background-color: #303030;
                    }              
                    QToolButton:hover {
                        background-color: #373737;
                }""")
            else:
                self.selected_button.setStyleSheet("QToolButton { background-color: none; } QToolButton:hover {background-color: #373737;}")    
        
        self.stacked_widget.setCurrentIndex(index)

    def show_page_chatbot(self, index): 
        if self.selected_button:
            if self.prev_session == self.session_id:
                    # print("yes")
                    self.selected_button.setStyleSheet(""" 
                    QToolButton {   
                        background-color: #303030;
                    }              
                    QToolButton:hover {
                        background-color: #373737;
                }""")
            else:
                self.selected_button.setStyleSheet("QToolButton { background-color: none; } QToolButton:hover {background-color: #373737;}")

        if index == 1:
            # print(self.input_text.toPlainText())
            self.reset_chatbot_page()
            self.start_chat(self.input_text.toPlainText())
            
        self.stacked_widget.setCurrentIndex(index)
    


    def toggle_navbar(self):
        self.dock_sidebar.setVisible(not self.dock_sidebar.isVisible())



    def create_summarizer_page(self):
        page = QWidget()
        layout = QVBoxLayout()
        top_layout = QHBoxLayout()
        top_layout.addWidget(self.navbar_button)
        
        
        # Title
        title = QLabel("TERMS AND CONDITIONS SUMMARIZER")
        title.setContentsMargins(0, 0, 0, 25)
        title.setObjectName("titleLabel")
        title.setAlignment(Qt.AlignCenter)
        title_spacer = QSpacerItem(20, 20, QSizePolicy.Expanding, QSizePolicy.Minimum)
    
        top_layout.addItem(title_spacer)
        top_layout.addWidget(title)
        top_layout.addStretch()        
        layout.addLayout(top_layout)

        # Text input
        self.input_text = QTextEdit()
        self.input_text.setFixedHeight(160)
        self.input_text.setPlaceholderText("Paste your terms and conditions text here or upload a file...")
        layout.addWidget(self.input_text)

        # Add summary type selection buttons with Basic as default selected
        self.summary_type_buttons = {
            "Basic": QPushButton("Basic"),
            "Specific": QPushButton("Specific"),
            "Abstractive": QPushButton("Abstractive"),
        }

        # Set button properties and select Basic by default
        summary_type_layout = QHBoxLayout()
        summary_type_layout.addStretch(1)  # Left alignment

        for key, button in self.summary_type_buttons.items():
            button.setCheckable(True)
            button.setCursor(Qt.PointingHandCursor)
            button.setStyleSheet("""
                QPushButton {
                    color : #4169E1;
                    min-height: 24px;
                    min-width: 180px;
                    background-color: #2b2b2b;    
                    border: 2px solid #4169E1; 
                    font-weight: bold;     
                    
                }
                QPushButton:checked {
                    background-color: #4169E1;
                    color : white;
                }
            """)
            button.clicked.connect(lambda checked, b=button: self.select_summary_type(b))
            summary_type_layout.addWidget(button)
            
        summary_type_layout.addStretch(1)  # Right alignment of summary type buttons

        # Add the attachment button at the far right
        attach_button = QPushButton()
        attach_button.setCursor(Qt.PointingHandCursor)
        attach_button.setIcon(QIcon(self.resource_path("./icons/attach_icon.png")))  # Replace with the path to your icon
        attach_button.setIconSize(QSize(20, 20))
        attach_button.setToolTip("Extract text from a file.")
        attach_button.setStyleSheet("""
            QPushButton {
                background-color: #3c3f41;
                border-radius: 18px;
                min-width: 24px;
                min-height: 24px;
            }
            QPushButton:hover {
                background-color: #505354;
            }
        """)
        attach_button.clicked.connect(self.load_attachment_summarizer)
        summary_type_layout.addWidget(attach_button)

        # Set Basic button as default selected
        self.summary_type_buttons["Basic"].setChecked(True)
        self.selected_summary_type = "Basic"

        layout.addLayout(summary_type_layout)


        # Summarize button and attachment button
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        summarize_button = QPushButton("Summarize")
        summarize_button.setCursor(Qt.PointingHandCursor)
        summarize_button.setStyleSheet("""
            QPushButton {
                min-width: 120px;
                min-height: 24px;
                margin-right: 43px;
            }
        """)
        
        summarize_button.clicked.connect(self.generate_summary)
        button_layout.addWidget(summarize_button)


        centered_button_layout = QHBoxLayout()
        centered_button_layout.addStretch(1)
        centered_button_layout.addLayout(button_layout)
        centered_button_layout.addStretch(1)
        layout.addLayout(centered_button_layout)
        layout.addSpacing(15)
        
        # Output box for summary
        self.output_text = QTextEdit()


        self.output_text.setStyleSheet("""
            QTextEdit {
            background-color: #2b2b2b;  /* Match background color */
            color: #ffffff;
            border: none;  /* Remove border */
            }                           
            
            QScrollArea {
                background-color: #2b2b2b;  /* Sidebar background */
                border: none;  /* Remove border */
            }
            QScrollBar:vertical {
                border: none;
                background-color: #2b2b2b;  /* Dark background for scrollbar */
                width: 10px;
                margin: 0px 0 0px 0;
                border-radius: 0px;
            }
            QScrollBar::handle:vertical {
                background-color: #5a5a5a;  /* Scroll handle color */
                min-height: 20px;
                border-radius: 0px;
            }
            QScrollBar::handle:vertical:hover {
                background-color: #6f6f6f;  /* Hover color for scroll handle */
            }
            
            QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical {
                background-color: #2b2b2b; 
            }
        """)

        self.output_text.setReadOnly(True)
        self.output_text.setPlaceholderText("Summary will appear here...")
        layout.addWidget(self.output_text)

        # Set layout for the page
        # page.setLayout(layout)

        # Create the bottom layout for the new buttons
        # Create the bottom layout for the new buttons
        self.bottom_button_layout = QHBoxLayout()
        self.bottom_button_layout.setContentsMargins(0, 5, 10, 0)

        # Ask Chatbot button (aligned to the left)
        self.ask_chatbot_button = QPushButton("Ask Chatbot")
        self.ask_chatbot_button.setCursor(Qt.PointingHandCursor)
        self.ask_chatbot_button.setFixedWidth(100)
        self.ask_chatbot_button.clicked.connect(lambda: self.show_page_chatbot(1))
        self.bottom_button_layout.addWidget(self.ask_chatbot_button)

        # Spacer to push other buttons to the right
        self.bottom_button_layout.addStretch()

        # Copy button
        copy_button = QPushButton()
        copy_button.setCursor(Qt.PointingHandCursor)
        copy_button.setIcon(QIcon(self.resource_path("./icons/copy_icon(3).png")))
        copy_button.setStyleSheet("""
            QPushButton {
                background-color:none;
                border : none;
                font-size: 24px;
            }
            QPushButton:hover {
                background-color: #505354;
            }
            
        """)
        copy_button.setToolTip("Copy the summary")
        copy_button.clicked.connect(lambda: QApplication.clipboard().setText(self.output_text.toPlainText()))
        self.bottom_button_layout.addWidget(copy_button)

        # Download button
        download_button = QPushButton()
        download_button.setCursor(Qt.PointingHandCursor)
        download_button.setToolTip("Download as a text file")
        download_button.setIcon(QIcon(self.resource_path("./icons/download_icon.png")))
        download_button.setStyleSheet("""
            QPushButton {
                background-color:none;
                border : none;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #505354;
            }
            
        """)
        download_button.clicked.connect(self.download_summary)
        self.bottom_button_layout.addWidget(download_button)

        # Set layout visibility to hidden initially
        self.bottom_button_layout_widget = QWidget()
        self.bottom_button_layout_widget.setLayout(self.bottom_button_layout)
        self.bottom_button_layout_widget.setVisible(False)  # Hide initially
        layout.addWidget(self.bottom_button_layout_widget)

        
        page.setLayout(layout)
        return page
    

    def select_summary_type(self, button):
        for key, b in self.summary_type_buttons.items():
            b.setChecked(False)
        button.setChecked(True)
        self.selected_summary_type = next(
            key for key, b in self.summary_type_buttons.items() if b == button
        )
        # print(f"Selected Summary Type: {self.selected_summary_type}")


    def download_summary(self):
        # Get the content of the output text box
        summary_text = self.output_text.toPlainText()
        
        if summary_text:
            # Open a file dialog to choose the save location
            file_path, _ = QFileDialog.getSaveFileName(self, "Save Summary", "output_summary.txt", "Text Files (*.txt)")
            
            # If a file path is chosen, save the summary text to the file
            if file_path:
                with open(file_path, 'w') as file:
                    file.write(summary_text)


    def download_summary_text(self, summary_text):
        if summary_text:
            # Open a file dialog to choose the save location
            file_path, _ = QFileDialog.getSaveFileName(self,"Save Summary", "output_summary.txt", "Text Files (*.txt)")
            
            # If a file path is chosen, save the summary text to the file
            if file_path:
                with open(file_path, 'w') as file:
                    file.write(summary_text)


    def load_attachment_summarizer(self):
        try:
            # extracted_text, path = self.extract_text_pytesseract()
            extracted_text, path = self.extract_text_easyocr()
        except Exception:
            return
        

        if extracted_text:
            self.input_text.setText(extracted_text)
        else:
            self.output_text.setText(f"Could not extract content from {path}")


    # def extract_text_pytesseract(self):
    #     file_dialog = QFileDialog()
    #     path, _ = file_dialog.getOpenFileName(self, "Open File", "", "Text files (*.txt *.pdf *.doc *.docx *.png *.jpg *.jpeg)")
    #     if path:
    #         extracted_text = ""
    #         if path.endswith(".txt"):
    #             with open(path, "r", encoding="utf-8") as file:
    #                 extracted_text = file.read()
    #         elif path.endswith(".docx") or path.endswith(".doc"):
    #             doc = Document(path)
    #             extracted_text = "\n".join([para.text for para in doc.paragraphs])
    #         elif path.endswith(".pdf"):
    #             pdf = fitz.open(path)
    #             for page_num in range(pdf.page_count):
    #                 page = pdf.load_page(page_num)
    #                 extracted_text += page.get_text()
    #             pdf.close()
    #         elif path.lower().endswith((".png", ".jpg", ".jpeg")):
    #             from PIL import Image
    #             import pytesseract
    #             pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
                
    #             image = Image.open(path)
    #             extracted_text = pytesseract.image_to_string(image)

    #     return (extracted_text,path)

    def extract_text_easyocr(self):
        file_dialog = QFileDialog()
        path, _ = file_dialog.getOpenFileName(self, "Open File", "", "Text files (*.txt *.pdf *.doc *.docx *.png *.jpg *.jpeg)")
        if path:
            extracted_text = ""
            if path.endswith(".txt"):
                with open(path, "r", encoding="utf-8") as file:
                    extracted_text = file.read()
            elif path.endswith(".docx") or path.endswith(".doc"):
                doc = Document(path)
                extracted_text = "\n".join([para.text for para in doc.paragraphs])
            elif path.endswith(".pdf"):
                pdf = fitz.open(path)
                for page_num in range(pdf.page_count):
                    page = pdf.load_page(page_num)
                    extracted_text += page.get_text()
                pdf.close()
            elif path.lower().endswith((".png", ".jpg", ".jpeg")):
                reader = Reader(['en'])  # Specify the language(s) as needed
                result = reader.readtext(path, detail=0)  # detail=0 for plain text output
                extracted_text = "\n".join(result)
            
        return (extracted_text,path)

    def generate_summary(self):
        self.bottom_button_layout_widget.setVisible(False)
        summary_type = getattr(self, 'selected_summary_type', "Basic")  # Default to Basic if not selected
        text = self.input_text.toPlainText()
        final_summary = ""

        if not text:
            self.output_text.setPlainText("Please enter the text or attach file.")
            return

        if self.selected_summary_type == "Basic":
            try:
                summary = st_generate_summary(text, reducing_factor= 8)  # Call the imported function
                final_summary = summary
            except Exception as e:
                self.output_text.setPlainText("Something went wrong! Execption: " + str(e))
                return

        elif self.selected_summary_type == "Specific":
            try: 
                summary = ag_generate_summary(text,num_sentences=4)  # Call the imported function
                for topic, sentences in summary.items():
                    final_summary += f"---- {topic} ----\n"
                    for sentence in sentences:
                        final_summary += f"- {sentence}\n"
                    final_summary += "\n"

            except Exception as e:
                self.output_text.setPlainText("Something went wrong! Execption: " + str(e))
                return

        elif self.selected_summary_type == "Abstractive":
            try:
                summary = generate_api_summary(text)  # Call the imported function
                if summary == "api_error":
                    self.output_text.setPlainText("API key not found. Please set the GROQ_API_KEY environment variable.")
                    return
                final_summary = summary
            except Exception as e:
                self.output_text.setPlainText("Something went wrong! Execption: " + str(e))
                return
            
        else:
            summary = "Invalid summary type selected."

        if len(final_summary) <= 1:
            self.output_text.setPlainText("Not enough data to generate summary.")
            return

        self.db.add_session_entry(self.session_id, summary_type, text, final_summary)
    
        self.output_text.setPlainText(final_summary)
        self.bottom_button_layout_widget.setVisible(True)
        QTimer.singleShot(1000, lambda: self.update_sidebar_sessions())



    def create_session_page(self, session_id, session_name):
        # Fetch and display session details
        summaries = self.db.get_session_summaries(session_id)

        page = QWidget()
        main_layout = QVBoxLayout(page)
        # main_layout.setContentsMargins(10,10,10,10)  
        
        # Scrollable area
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        # scroll_area.setStyleSheet("QScrollArea { background-color: #2b2b2b; border: none; }")
        scroll_area.setStyleSheet("""
             QScrollArea {
                background-color: #212121;  /* Sidebar background */
                border: none;  /* Remove border */
            }
            QScrollBar:vertical {
                border: none;
                width: 4px;
                border-radius: 0px;
                margin: 0px;
            }
            QScrollBar::handle:vertical {
                background-color: #3d3d3d;  /* Scroll handle color */
                min-height: 20px;
                border-radius: 0px;
            }
            QScrollBar::handle:vertical:hover {
                background-color: #6f6f6f;  /* Hover color for scroll handle */
            }
            
            QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical {
                background-color:#212121; 
            }
        """)

        # Inner widget for scroll area content
        content_widget = QWidget()
        layout = QVBoxLayout(content_widget)
        layout.setSpacing(15)  # Add spacing between sections for visual clarity
        layout.setContentsMargins(0,0,45,0)

        # Top layout with title
        top_layout = QHBoxLayout()
        top_layout.setContentsMargins(0,0,0,13)
        
        session_alias = self.db.get_session_alias(session_id)
        if session_alias == "None" or session_alias == "Session started":
            title = QLineEdit(f"Session {''.join(session_name.split()[:-2])}")
        else:
            title = QLineEdit(session_alias)
        
        title.setObjectName("titleEdit")
        title.setReadOnly(True)  # Initially set to read-only
        title.setCursorPosition(0)
        title.setContentsMargins(0, 0, 0, 0)
        title.setStyleSheet("""
            QLineEdit {
                border: none;
                color: white;
                font-size: 28px;
            }
           
        """)
 
        top_layout.addWidget(title)
        top_layout.addStretch()
        
        # title_spacer = QSpacerItem(0, 0, QSizePolicy.Expanding, QSizePolicy.Minimum)
        

        options_button = QToolButton()
        options_button.setCursor(Qt.PointingHandCursor)
        options_button.setIcon(QIcon(self.resource_path("./icons/threedots_icon.png")))  # Replace with path to your triple dot icon
        options_button.setIconSize(QSize(18, 18))
        options_button.setStyleSheet("""QToolButton { background-color : none; 
                                                    padding: 5px;
                                                    
                                                    border-radius: 14px;
                                                    border: none} 
                                        QToolButton:hover {background-color: #373737;}
                                        QToolButton::menu-indicator{ width:0px; }; 
                                    """)
        top_layout.addWidget(options_button, alignment=Qt.AlignRight | Qt.AlignVCenter)
        
        # Setup for the options menu
        options_menu = QMenu()
        options_menu.setStyleSheet("""
            QMenu {
                border: 1px solid #444444;
                background-color: #2b2b2b;
                border-radius: 8px;
                padding: 5px;
            }
            QMenu::item {
                padding: 8px 20px;
                color: #f0f0f0;
                font-size: 16px;
            }
            QMenu::item:selected {
                background-color: #575757;
                border-radius: 8px;
            }
            QMenu::icon {
                padding: 8px;
            }
        """)
        # Actions with icons
        rename_action = QAction(QIcon(self.resource_path("./icons/rename_icon.png")), "Rename Session", self)
        share_action = QAction(QIcon(self.resource_path("./icons/export_icon.png")), "Export Session", self)
        delete_action = QAction(QIcon(self.resource_path("./icons/delete_icon.png")), "Delete Session", self)
        

        # Connect the menu actions to functions (leave functions unimplemented as requested)
        rename_action.triggered.connect(lambda: self.enable_title_edit(title, session_id))
        share_action.triggered.connect(lambda checked, idx=(session_id, summaries): self.export_session(idx))
        delete_action.triggered.connect(lambda checked, idx=session_id: self.delete_session(idx))

        # Add actions to the menu()
        options_menu.addAction(rename_action)
        options_menu.addAction(share_action)
        options_menu.addAction(delete_action)

        # Attach the menu to the options button
        options_button.setMenu(options_menu)
        options_button.setPopupMode(QToolButton.InstantPopup)

        # top_layout.addStretch()
        layout.addLayout(top_layout)

        
        
        # Display summaries (non-editable)
        for summary in reversed(summaries):
            layout.addSpacing(10)

            time_label = QLabel(f"{summary[2]:}")
            time_label.setStyleSheet("""
                font-size: 20px;
                font-weight: bold;
                color: white;
                padding-bottom: -10px;
            """)
            layout.addWidget(time_label)

            line_separator = QFrame()
            line_separator.setFrameShape(QFrame.HLine)
            line_separator.setFrameShadow(QFrame.Sunken)
            line_separator.setStyleSheet("color: #5a5a5a; border-top: 5px solid #5a5a5a;")
            layout.addWidget(line_separator)

            # Input text
            
            
            input_text_label = QLabel("Input Text:")
            input_text_label.setStyleSheet("font-weight: bold;")
            layout.addWidget(input_text_label)
            
            input_text_display = QTextEdit()
            input_text_display.setPlainText(summary[4])  # Set text with line breaks preserved
            input_text_display.setReadOnly(True)
            input_text_display.setMinimumHeight(120)
            input_text_display.setMaximumHeight(200)
            input_text_display.setStyleSheet(self.input_text.styleSheet())
            layout.addWidget(input_text_display)

            layout.addSpacing(15)

            # Summary text
            summary_label = QLabel(f"{summary[3]} Summary:")  # Use summary type as label
            summary_label.setStyleSheet("font-weight: bold;")
            layout.addWidget(summary_label)
            
            summary_text_display = QTextEdit()
            summary_text_display.setPlainText(summary[5])  # Set text with line breaks preserved
            summary_text_display.setReadOnly(True)
            summary_text_display.setMinimumHeight(300)
            summary_text_display.setMaximumHeight(600)
            summary_text_display.setStyleSheet(self.output_text.styleSheet())
            layout.addWidget(summary_text_display)

            # Bottom button layout for copy and download
            bottom_button_layout = QHBoxLayout()
            bottom_button_layout.addStretch()

            # Copy button
            copy_button = QPushButton()
            copy_button.setCursor(Qt.PointingHandCursor)
            copy_button.setIcon(QIcon(self.resource_path("./icons/copy_icon(3).png")))
            copy_button.setStyleSheet("""
                QPushButton {
                    background-color: none;
                    border: none;
                    font-size: 24px;
                }
                QPushButton:hover {
                    background-color: #505354;
                }
            """)
            copy_button.setToolTip("Copy the summary")
            # Connect the copy button to copy the specific summary's text
            copy_button.clicked.connect(lambda _, text=summary[5]: QApplication.clipboard().setText(text))
            bottom_button_layout.addWidget(copy_button)

            # Download button
            download_button = QPushButton()
            download_button.setCursor(Qt.PointingHandCursor)
            download_button.setToolTip("Download as a text file")
            download_button.setIcon(QIcon(self.resource_path("./icons/download_icon.png")))
            download_button.setStyleSheet("""
                QPushButton {
                    background-color: none;
                    border: none;
                    font-size: 12px;
                }
                QPushButton:hover {
                    background-color: #505354;
                }
            """)
            # Connect the download button to download the specific summary
            # Download button
            download_button.clicked.connect(lambda _, text=summary[5]: self.download_summary_text(text))
            bottom_button_layout.addWidget(download_button)

            # Add the button layout to the main layout
            bottom_button_layout_widget = QWidget()
            bottom_button_layout_widget.setLayout(bottom_button_layout)
            layout.addWidget(bottom_button_layout_widget)
            
        # Set layout and scroll area
        scroll_area.setWidget(content_widget)
        main_layout.addWidget(scroll_area)
        
        return page


    def enable_title_edit(self, title,session_id):
        title.setReadOnly(False)
        title.setFocus()
        title.selectAll()  # Select existing text for easier editing

        # Handle renaming when user clicks away or presses Enter
        def finish_renaming():
            new_name = title.text().strip()
            title.setReadOnly(True)  # Set back to read-only
            title.setText(new_name)
            self.db.update_session_name(session_id, new_name)  # Update in the database
            title.clearFocus()  # Remove focus to finish editing mode
            title.setCursorPosition(0)

        title.editingFinished.connect(finish_renaming)
        title.returnPressed.connect(finish_renaming)
        

    def show_session(self, session):

        if self.selected_button:
            if self.prev_session == self.session_id:
                    # print("yes")
                    self.selected_button.setStyleSheet(""" 
                    QToolButton {   
                        background-color: #303030;
                    }              
                    QToolButton:hover {
                        background-color: #373737;
                }""")
            else: 
                self.selected_button.setStyleSheet("QToolButton { background-color: none; } QToolButton:hover {background-color: #373737;}")

        button = session[2]
        button.setStyleSheet("QToolButton { background-color: #4169E1; }")

        self.selected_button = button

        session_id = session[0]
        self.prev_session = session_id
        session_name = session[1]


        # Create the session page
        session_page = self.create_session_page(session_id, session_name)

        # Add the session page to the stacked widget and display it
        self.stacked_widget.addWidget(session_page)
        self.stacked_widget.setCurrentWidget(session_page)


    def export_session(self, arg):
        summaries = reversed(arg[1])
        session_id = arg[0]
        session = self.db.get_log_entry(session_id)
        session_name = session[1] + " " +session[2]

        # Open a file dialog to select the save location and file name
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getSaveFileName(
            None,
            "Save PDF",
            f"session_{session[1]} {session_id}_summary.pdf",
            "PDF Files (*.pdf);;All Files (*)",
            options=options
        )

        # Proceed if a file path was selected
        if file_path:
            # Create the PDF document
            doc = SimpleDocTemplate(file_path, pagesize=A4)
            elements = []
            styles = getSampleStyleSheet()
            

            # Title for the PDF
            title = Paragraph(f"Session {session_name} - Summaries", styles["Title"])
            elements.append(title)
            elements.append(Spacer(1, 0.5 * inch))

            # Loop through summaries and add each summary to the PDF
            for idx, summary in enumerate(summaries, start=1):
                # Extract details from the summary tuple
                _, session_id, time, summary_type, input_text, summary_text = summary

                # Add summary number and type
                elements.append(Paragraph(f"Summary {idx} - {summary_type}", styles["Heading2"]))
                elements.append(Spacer(1, 0.2 * inch))

                # Define a style with deep blue color for summaries
                deep_blue_style = styles["BodyText"].clone('DeepBlueBodyText')
                deep_blue_style.textColor = colors.HexColor("#003366")  # Deep blue color

                formatted_input_text = input_text.replace('\n', '<br />')  # Replace newlines with <br /> tags
                elements.append(Paragraph("Input Text:", styles["Heading3"]))
                elements.append(Paragraph(formatted_input_text, deep_blue_style))
                elements.append(Spacer(1, 0.2 * inch))

                # Add summary text section with default color, preserving line breaks
                formatted_summary_text = summary_text.replace('\n', '<br />')  # Replace newlines with <br /> tags
                elements.append(Paragraph(f"{summary_type} Summary:", styles["Heading3"]))
                elements.append(Paragraph(formatted_summary_text, styles["BodyText"]))
                elements.append(Spacer(1, 0.5 * inch))

                # Divider line between summaries
                divider_table = Table([[""]], colWidths=[7.5 * inch])
                divider_table.setStyle(TableStyle([('LINEABOVE', (0, 0), (-1, -1), 1, colors.grey)]))
                elements.append(divider_table)
                elements.append(Spacer(1, 0.2 * inch))

            # Build the PDF
            doc.build(elements)

            # print(f"PDF generated successfully: {file_path}")
        else:
            pass
            # print("PDF export cancelled.")


    def delete_session(self,session_id):
        session = self.db.get_log_entry(session_id)
        # print(session_id)
        # Create a new page to display the deletion message
        delete_page = QWidget()
        delete_layout = QVBoxLayout(delete_page)
        delete_layout.setContentsMargins(20, 20, 20, 20)
        
        # Set background color and styling for the delete page
        delete_page.setStyleSheet("""
            QWidget {
                background-color: none;
            }
            QLabel#deleteMessage {
                font-size: 24px;
                font-weight: bold;
                color: white;
                text-align: center;
            }
        """)

        # Create and style the deletion message label
        delete_message = QLabel(f"The Session {session[1]} is deleted.")
        delete_message.setAlignment(Qt.AlignCenter)
        delete_message.setObjectName("deleteMessage")
        
        # Add the message label to the layout
        delete_layout.addStretch()
        delete_layout.addWidget(delete_message)
        delete_layout.addStretch()

        self.db.reset_summary_generated(session_id)
        
        # Replace the current page with the delete page in the main layout
        self.stacked_widget.addWidget(delete_page)
        self.stacked_widget.setCurrentWidget(delete_page)
        QTimer.singleShot(500, lambda: self.update_sidebar_sessions())




    def create_chatbot_page(self):
        page = QWidget()

        self.apply_dark_theme()

        # Groq Client Setup
        api_key = os.getenv("GROQ_API_KEY")
        if not api_key:
            main_layout = QVBoxLayout(page)
            missing_api_key = QLabel("API key not found. Please set the GROQ_API_KEY environment variable.")
            missing_api_key.setAlignment(Qt.AlignCenter)
            main_layout.addWidget(missing_api_key)
            return page
            
        self.client = Groq(api_key=api_key)

        # Session state
        self.chat_history = []
        self.max_tokens = 5000
        
        # main_widget = QWidget(page)
        main_layout = QVBoxLayout(page)


        self.header_container = QWidget()
        self.header_container.setFixedHeight(400)  # Adjust as needed
        header_layout = QVBoxLayout(self.header_container)
        header_layout.setContentsMargins(0, 0, 0, 0)
        header_layout.setAlignment(Qt.AlignTop)

        self.chatbot_title = QLabel("TERMS AND CONDITIONS CHATBOT")
        self.chatbot_title.setContentsMargins(0, 0, 0, 45)
        self.chatbot_title.setObjectName("titleLabel")
        self.chatbot_title.setAlignment(Qt.AlignCenter)
        # title.setStyleSheet("""
        #     QLabel#titleLabel {
        #         font-size: 24px;
        #         font-weight: bold;
        #         color: #ffffff;
        #     }
        # """)
        header_layout.addWidget(self.chatbot_title)

        # Input Text label
        self.chatbot_input_label = QLabel("Input Text:")
        self.chatbot_input_label.setVisible(False)
        self.chatbot_input_label.setContentsMargins(3, 0, 0, 2)  # Adjust margins for proper alignment
        self.chatbot_input_label.setStyleSheet("""
            QLabel {
                font-size: 14px;
                color: #CCCCCC;
            }
        """)
        header_layout.addWidget(self.chatbot_input_label, alignment=Qt.AlignLeft)

        # Input Text Area
        self.chatbot_input_text = QTextEdit()
        self.chatbot_input_text.setFixedHeight(250)
        self.chatbot_input_text.setPlaceholderText("Paste your terms and conditions text here or upload a file...")
        header_layout.addWidget(self.chatbot_input_text)

        # Buttons Layout
        button_container = QWidget()
        button_layout = QHBoxLayout(button_container)
        button_layout.setContentsMargins(0, 0, 0, 0)

        button_layout.addStretch()

        # Attach Button
        self.chatbot_attach_button = QPushButton()
        self.chatbot_attach_button.setCursor(Qt.PointingHandCursor)
        self.chatbot_attach_button.setIcon(QIcon(self.resource_path("./icons/attach_icon.png")))  # Replace with the path to your icon
        self.chatbot_attach_button.setIconSize(QSize(20, 20))
        self.chatbot_attach_button.setToolTip("Extract text from a file.")
        self.chatbot_attach_button.setStyleSheet("""
            QPushButton {
                background-color: #3c3f41;
                border-radius: 18px;
                min-width: 24px;
                min-height: 24px;
            }
            QPushButton:hover {
                background-color: #505354;
            }
        """)
    
        self.chatbot_attach_button.clicked.connect(lambda checked, param=(self.chatbot_input_text): self.load_attachment_chatbot(param))

        

        # Summarize Button
        self.chatbot_chat_button = QPushButton("Chat")
        self.chatbot_chat_button.setCursor(Qt.PointingHandCursor)
        self.chatbot_chat_button.setStyleSheet("""
            QPushButton {
                min-width: 120px;
                min-height: 24px;
                
            }
        """)
        # summarize_button.clicked.connect(self.toggle_chat_visibility)

        button_layout.addWidget(self.chatbot_chat_button, alignment=Qt.AlignCenter)
        button_layout.addStretch()

        button_layout.addWidget(self.chatbot_attach_button, alignment=Qt.AlignRight)
        
        # button_layout.addWidget(summarize_button)
        # button_layout.addWidget(attach_button)
        # button_layout.addStretch()

        header_layout.addWidget(button_container)

        # Add Header Section to Main Layout
        main_layout.insertWidget(0, self.header_container, stretch=0)



        # Scrollable Chat Display
        self.chatbot_scroll_area = QScrollArea()
        self.chatbot_scroll_area.setVisible(False)
        self.chatbot_scroll_area.setWidgetResizable(True)
        # self.scroll_area.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.chat_container = QWidget()
        self.chat_layout = QVBoxLayout(self.chat_container)
        self.chat_layout.setSpacing(0)
        self.chat_layout.addStretch()
        self.chatbot_scroll_area.setWidget(self.chat_container)

        # Customize Scroll Bar
        self.chatbot_scroll_area.setStyleSheet("""
            QScrollArea {
                background-color: #222222;  /* Sidebar background */
                border: none;  /* Remove border */
            }
            QScrollBar:vertical {
                border: none;
                background-color: #2b2b2b;  /* Dark background for scrollbar */
                width: 5px;
                margin: 0px 0 0px 0;
                border-radius: 0px;
            }
            QScrollBar::handle:vertical {
                background-color: #3f3f3f;  /* Scroll handle color */
                min-height: 20px;
                border-radius: 0px;
            }
            QScrollBar::handle:vertical:hover {
                background-color: #4f4f4f;  /* Hover color for scroll handle */
            }
            QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical {
                background-color:#222222; 
            }
        """)

        # Input Area
        self.chatbot_input_container = QWidget()
        # input_container.setFixedWidth(1000)
        # input_container.setFixedSize(700, 70)
        self.chatbot_input_container.setFixedHeight(60)
        self.chatbot_input_container.setMinimumWidth(700)
        self.chatbot_input_container.setMaximumWidth(1200)
        
        # input_container.setContentsMargins(10,0,10,0)
        # input_container.setFixedHeight(70)
        self.chatbot_input_container.setStyleSheet(
            """
            background-color: #383838; 
            border-radius: 28px; 
            padding: 0px;
            """
        )

        input_layout = QHBoxLayout(self.chatbot_input_container)
        input_layout.setContentsMargins(10, 0, 10, 0)
        input_layout.setSpacing(0)


        # Reset Button
        self.reset_button = QPushButton()
        self.reset_button.setIcon(QIcon(self.resource_path("./icons/reset_icon(3).png")))  # Replace with path to your reset icon
        self.reset_button.setIconSize(QSize(28, 28))
        self.reset_button.setFixedHeight(50)
        self.reset_button.setToolTip("Reset Conversation")
        self.reset_button.setStyleSheet(
            """
            QPushButton {
                background-color: none; 
                color: white;
                border: none; 
                padding-left: 6px;  
                padding-right: 10px; 
            }
        
            """
        )
        self.reset_button.setCursor(Qt.PointingHandCursor)  # Change cursor to pointer on hover
        
        input_layout.addWidget(self.reset_button)

        self.chat_input = QTextEdit()
        self.chat_input.setFixedSize(950, 50)
        self.chat_input.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)  # Hide vertical scroll bar
        self.chat_input.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.chat_input.setStyleSheet(
            """
            background-color: transparent; 
            color: #ffffff; 
            border: none; 
            font-size: 16px; 
            padding: 2px;
            """
        )
        self.chat_input.setPlaceholderText("Type a message...")
        self.chat_input.setFixedHeight(50)  # Restrict max height to around 6 lines
        self.chat_input.installEventFilter(self)
        input_layout.addWidget(self.chat_input)

        self.send_button = QPushButton()
        self.send_button.setIcon(QIcon(self.resource_path("./icons/send_icon_hover.png")))  # Replace with path to your menu icon
        self.send_button.setIconSize(QSize(38, 38))

        self.send_button.setFixedHeight(50)
        self.send_button.setStyleSheet(
            """
            QPushButton {
                background-color: transparent; 
                border: none; 
            }
            QPushButton:hover {
                background-color: transparent;
            }
            """
            )
        self.send_button.setCursor(Qt.PointingHandCursor)  # Change cursor to pointer on hover
        self.send_button.clicked.connect(self.send_message)
        input_layout.addWidget(self.send_button)

        main_layout.addWidget(self.chatbot_scroll_area, stretch=8)  # Allocate most space to the scroll area
        main_layout.addWidget(self.chatbot_input_container, stretch=1, alignment=Qt.AlignHCenter)

        self.chatbot_scroll_area.setVisible(False)
        self.chatbot_input_container.setVisible(False)
        self.chatbot_chat_button.clicked.connect(self.start_chat)
        self.reset_button.clicked.connect( self.reset_chatbot_page)
        # main_layout.addStretch(1)
        main_layout.setContentsMargins(20,20,20,20)

        # self.setCentralWidget(main_widget)
        self.setCentralWidget(page)

        return page
    

    def reset_chatbot_page(self):
        self.chat_history = []
    
        # Clear chat container
        for i in reversed(range(self.chat_layout.count())):
            widget = self.chat_layout.itemAt(i).widget()
            if widget:
                widget.deleteLater()
        
        # Reset input fields
        self.chat_input.clear()
        self.chat_input.setPlaceholderText("Type a message...")
        
        self.chatbot_title.setVisible(True)
        self.header_container.setFixedHeight(400) 
        self.chatbot_input_label.setVisible(False)
        self.chatbot_input_text.setFixedHeight(250)
        self.chatbot_input_text.setReadOnly(False)
        self.chatbot_scroll_area.setVisible(False)
        self.chatbot_input_container.setVisible(False)
        self.chatbot_chat_button.setVisible(True)
        self.chatbot_attach_button.setVisible(True)


    def start_chat(self,summary = None):
        

        if not summary :
            user_input = self.chatbot_input_text.toPlainText().strip()
            if not user_input or user_input == "Please provide the input text.":
                self.chatbot_input_text.setText("Please provide the input text.")
                return
        else:
            # print(summary)
            user_input = summary
            self.chatbot_input_text.setText(summary)
        self.toggle_chat_visibility()
        
        user_input = generate_api_summary(user_input)
        # print(user_input)
        system_message = {
            "role": "system",
            "content": (
                "You must strictly limit the conversation to the provided terms and conditions. "
                "Do not answer question that are absolutely unrelated to the provided content "
                "Talk point to point. "
                f"Here are the terms and conditions:\n {user_input}"
            )
        }
        
        self.chat_history.append(system_message)
    

    def toggle_chat_visibility(self):
        # is_summarized = self.input_text.toPlainText().strip() != ""  # Check if input text is not empty
        self.chatbot_title.setVisible(False)
        self.header_container.setFixedHeight(150) 
        self.chatbot_input_label.setVisible(True)
        self.chatbot_input_text.setFixedHeight(120)
        self.chatbot_input_text.setReadOnly(True)
        self.chatbot_scroll_area.setVisible(True)
        self.chatbot_input_container.setVisible(True)
        self.chatbot_chat_button.setVisible(False)
        self.chatbot_attach_button.setVisible(False)


    def load_attachment_chatbot(self,param):
        input_text = param
        extracted_text, path = self.extract_text_pytesseract()
        # extracted_text, path = self.extract_text_easyocr()

        if extracted_text:
            input_text.setText(extracted_text)
        else:
            input_text.setText(f"Could not extract content from {path}")
        

    def apply_dark_theme(self):
        """Apply a dark theme to the application."""
        dark_palette = QPalette()
        dark_palette.setColor(QPalette.Window, QColor(30, 30, 30))
        dark_palette.setColor(QPalette.WindowText, QColor(255, 255, 255))
        dark_palette.setColor(QPalette.Base, QColor(45, 45, 45))
        dark_palette.setColor(QPalette.AlternateBase, QColor(30, 30, 30))
        dark_palette.setColor(QPalette.ToolTipBase, QColor(255, 255, 255))
        dark_palette.setColor(QPalette.ToolTipText, QColor(255, 255, 255))
        dark_palette.setColor(QPalette.Text, QColor(255, 255, 255))
        dark_palette.setColor(QPalette.Button, QColor(45, 45, 45))
        dark_palette.setColor(QPalette.ButtonText, QColor(255, 255, 255))
        dark_palette.setColor(QPalette.Highlight, QColor(70, 130, 180))
        dark_palette.setColor(QPalette.HighlightedText, QColor(255, 255, 255))
        QApplication.setPalette(dark_palette)


    def send_message(self):
        user_message = self.chat_input.toPlainText().strip()
        if not user_message:
            return
        self.chat_input.clear()

        # Display user message
        self.add_message(user_message, "user")
        self.chat_history.append({"role": "user", "content": user_message})

        # Fetch assistant response
        self.fetch_response()
 
 
    def fetch_response(self):
        try:
            # API Call
            response_chunks = self.client.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=self.chat_history,
                max_tokens=self.max_tokens,
                
            )
            # print("API Response:", response_chunks)
                
            assistant_response = self.generate_chat_responses(response_chunks)
            self.chat_history.append({"role": "assistant", "content": assistant_response})
            self.add_message(assistant_response, "assistant")
        except Exception as e:
            self.add_message(f"Error: {str(e)}", "assistant")


    def generate_chat_responses(self, chat_completion):
        try:
            content = chat_completion.choices[0].message.content
            return content.strip() if content else ""
        except AttributeError:
            # Log the error for debugging
            print("Unexpected response structure:", chat_completion)
            raise ValueError("Invalid response format from the API.")

   
    def add_message(self, message, sender):
        """Add a message bubble to the chat display."""
        bubble = QLabel(message)
        bubble.setWordWrap(True)
        bubble.setStyleSheet(
            "border-radius: 15px; padding: 8px;"
            f"background-color: {'#005f73' if sender == 'user' else 'none'};"
            "color: #ffffff;  font-size: 16px;"  # Increased font size and max width
        )
        bubble.setAlignment(Qt.AlignLeft if sender == "assistant" else Qt.AlignRight)
        bubble.setTextInteractionFlags(Qt.TextSelectableByMouse | Qt.TextSelectableByKeyboard)
        bubble.setCursor(Qt.IBeamCursor)

        bubble_layout = QHBoxLayout()
        bubble_layout.setContentsMargins(0,0,8,8)
        if sender == "assistant":
            bubble_layout.addWidget(bubble)
            # bubble_layout.addStretch()
        else:
            bubble_layout.addStretch()
            bubble_layout.addWidget(bubble)
            

        bubble_frame = QFrame()
        bubble_frame.setLayout(bubble_layout)
        self.chat_layout.insertWidget(self.chat_layout.count() - 1, bubble_frame)

        # Auto-scroll to the bottom
        
        QTimer.singleShot(20, self.scroll_to_bottom)
  
   
    def scroll_to_bottom(self):
        scroll_bar = self.chatbot_scroll_area.verticalScrollBar()
        scroll_bar.setValue(scroll_bar.maximum())


       
    def eventFilter(self, source, event):
        if source == self.chat_input and event.type() == QEvent.KeyPress:
            if event.key() == Qt.Key_Return and not (event.modifiers() & Qt.ShiftModifier):
                self.send_message()
                return True  # Prevent further processing of the key event
        return super().eventFilter(source, event)



    def keyPressEvent(self, event):
        if event.key() == Qt.Key_F11:  # When F11 is pressed
            if self.isFullScreen():  # If the window is already full-screen, minimize
                self.showNormal()
            else:  # If the window is not full-screen, maximize it
                self.showFullScreen()

        elif event.key() == Qt.Key_Escape:  # When Esc is pressed
            if self.isFullScreen():  # If the window is full-screen, revert to normal
                self.showNormal()
        elif event.key() == Qt.Key_F5:  # When F5 is pressed
            QCoreApplication.quit()  # Quit the current application
            QProcess.startDetached(sys.executable, sys.argv)
        elif event.key() == Qt.Key_F1:  # Toggle sidebar visibility on F1
            self.dock_sidebar.setVisible(not self.dock_sidebar.isVisible())

    def resource_path(self,relative_path):
        """Get the absolute path to a resource, works for dev and PyInstaller."""
        if hasattr(sys, '_MEIPASS'):
            # PyInstaller sets _MEIPASS to the temp directory where resources are extracted
            return os.path.join(sys._MEIPASS, relative_path)
        return os.path.join(os.path.abspath("."), relative_path)
        
    def resizeEvent(self, event):
        self.close_button.setVisible( self.isFullScreen())

    def closeEvent(self, event):
        self.db.close()
        # Accept the event (allows the app to close)
        event.accept()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = TermsSummarizerApp()
    window.show()
    sys.exit(app.exec_())
